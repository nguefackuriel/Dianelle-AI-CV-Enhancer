"""
CV Exporter Module
Generate ATS-friendly DOCX and PDF files from CV content.
Provides multiple professional templates optimized for ATS parsing.
"""

import io
import re
from typing import Dict, List, Any, Optional
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE


class CVExporter:
    """Export optimized CVs to DOCX and PDF formats."""

    # ATS-safe fonts
    FONTS = {
        'classic': 'Calibri',
        'modern': 'Arial',
        'minimal': 'Garamond',
    }

    # Template color schemes (accent colors)
    COLORS = {
        'classic': RGBColor(0x1F, 0x47, 0x88),   # Navy blue
        'modern': RGBColor(0x2C, 0x3E, 0x50),     # Dark slate
        'minimal': RGBColor(0x33, 0x33, 0x33),     # Dark gray
        'professional': RGBColor(0x00, 0x66, 0x99), # Teal blue
        'executive': RGBColor(0x8B, 0x00, 0x00),    # Dark red
    }

    # Section order for ATS
    SECTION_ORDER = [
        'contact', 'summary', 'experience', 'education',
        'skills', 'certifications', 'projects', 'awards',
        'publications', 'volunteer', 'languages',
    ]

    def __init__(self):
        pass

    # ------------------------------------------------------------------
    # DOCX Generation
    # ------------------------------------------------------------------

    def generate_docx(
        self,
        cv_data: Dict[str, Any],
        template: str = 'classic',
        accent_color: str = 'classic',
        font_size: int = 11,
        include_sections: List[str] = None,
    ) -> io.BytesIO:
        """
        Generate an ATS-friendly DOCX file.

        Args:
            cv_data: Dict with 'name', 'contact', 'sections' (ordered dict of section_name -> content)
            template: Template style ('classic', 'modern', 'minimal')
            accent_color: Color scheme name
            font_size: Base font size
            include_sections: Which sections to include (None = all)

        Returns:
            BytesIO buffer containing the DOCX file
        """
        doc = Document()
        font_name = self.FONTS.get(template, 'Calibri')
        color = self.COLORS.get(accent_color, self.COLORS['classic'])

        # Configure default styles
        self._configure_styles(doc, font_name, font_size, color)

        # Set narrow margins for more content space
        for section in doc.sections:
            section.top_margin = Cm(1.5)
            section.bottom_margin = Cm(1.5)
            section.left_margin = Cm(2.0)
            section.right_margin = Cm(2.0)

        # Add contact/name header
        self._add_header(doc, cv_data, font_name, color, template)

        # Add sections in ATS-preferred order
        sections = cv_data.get('sections', {})
        ordered_sections = self._order_sections(sections, include_sections)

        for section_name, content in ordered_sections.items():
            if content and content.strip():
                self._add_section(doc, section_name, content, font_name, font_size, color, template)

        # Save to buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

    def generate_docx_from_text(
        self,
        cv_text: str,
        parsed_sections: Dict[str, str] = None,
        accepted_rewrites: Dict[str, str] = None,
        template: str = 'classic',
        contact_info: Dict[str, str] = None,
    ) -> io.BytesIO:
        """
        Convenience method to generate DOCX from raw text and optional rewrites.

        Args:
            cv_text: Original CV text
            parsed_sections: Parsed sections from ATS simulator
            accepted_rewrites: User-accepted rewrites per section
            template: Template style
            contact_info: Dict with 'name', 'email', 'phone', 'location', 'linkedin'
        """
        # Merge sections with rewrites
        sections = dict(parsed_sections or {})

        if accepted_rewrites:
            for section_name, content in accepted_rewrites.items():
                sections[section_name] = content

        # Extract contact info from header if not provided
        if not contact_info:
            contact_info = self._extract_contact_info(
                sections.get('header', '') or cv_text[:500]
            )

        # Remove header from sections (it's handled separately)
        sections.pop('header', None)
        sections.pop('contact', None)

        cv_data = {
            'name': contact_info.get('name', 'Your Name'),
            'contact': contact_info,
            'sections': sections,
        }

        return self.generate_docx(cv_data, template=template)

    # ------------------------------------------------------------------
    # PDF Generation (using fpdf2 or plain text fallback)
    # ------------------------------------------------------------------

    def _sanitize_pdf_text(self, text: str) -> str:
        """Sanitize Unicode characters to their safe Latin-1 equivalents for PDF generation."""
        replacements = {
            '\u2013': '-',    # en dash
            '\u2014': '-',    # em dash
            '\u201c': '"',    # curly double open
            '\u201d': '"',    # curly double close
            '\u2018': "'",    # curly single open
            '\u2019': "'",    # curly single close
            '\u201b': "'",
            '\u201a': "'",
            '\u201e': '"',
            '\u201f': '"',
            '\u2022': '·',    # bullet -> middle dot
            '\u2026': '...',  # ellipsis
            '\xa0': ' ',       # non-breaking space
            '\u200b': '',      # zero-width space
            '\u25aa': '·',     # small black square
            '\u25cf': '·',     # black circle
            '\u25cb': '·',     # white circle
            '\u25c6': '·',     # black diamond
            '\u25b8': '·',     # black right-pointing small triangle
            '\u25ba': '·',     # black right-pointing pointer
            '\u2192': '->',    # right arrow
            '\u00bb': '>>',    # right angle quote
            '•': '·',          # regular bullet
            '▪': '·',          # square bullet
            '◦': '·',          # circle bullet
            '·': '·',
        }
        for orig, rep in replacements.items():
            text = text.replace(orig, rep)
        return text

    def generate_pdf(
        self,
        cv_data: Dict[str, Any],
        template: str = 'classic',
    ) -> Optional[io.BytesIO]:
        """
        Generate a PDF file from CV data.
        Uses fpdf2 if available, otherwise returns None.
        """
        try:
            from fpdf import FPDF

            pdf = FPDF()
            pdf.set_auto_page_break(auto=True, margin=15)
            pdf.add_page()

            # Accent colors for PDF styling
            accent_rgbs = {
                'classic': (31, 71, 136),    # Navy Blue
                'modern': (44, 62, 80),      # Dark Slate
                'minimal': (51, 51, 51),     # Dark Gray
            }
            accent_rgb = accent_rgbs.get(template, (31, 71, 136))

            # Header (Name)
            pdf.set_font('Helvetica', 'B', 18)
            pdf.set_text_color(*accent_rgb)
            name = cv_data.get('name', 'Your Name')
            pdf.cell(0, 12, name, new_x="LMARGIN", new_y="NEXT", align='C')

            # Contact info
            contact = cv_data.get('contact', {})
            contact_line = ' | '.join(
                v for v in [
                    contact.get('email', ''),
                    contact.get('phone', ''),
                    contact.get('location', ''),
                    contact.get('linkedin', ''),
                ] if v
            )
            if contact_line:
                pdf.set_font('Helvetica', '', 9)
                pdf.set_text_color(100, 100, 100)
                pdf.cell(0, 6, contact_line, new_x="LMARGIN", new_y="NEXT", align='C')

            pdf.ln(5)

            # Sections
            sections = cv_data.get('sections', {})
            ordered = self._order_sections(sections)

            for section_name, content in ordered.items():
                if not content or not content.strip():
                    continue

                # Section header
                pdf.set_font('Helvetica', 'B', 12)
                pdf.set_text_color(*accent_rgb)
                header_text = section_name.replace('_', ' ').upper()
                pdf.cell(0, 8, header_text, new_x="LMARGIN", new_y="NEXT")

                # Horizontal line
                pdf.set_draw_color(*accent_rgb)
                pdf.line(10, pdf.get_y(), 200, pdf.get_y())
                pdf.ln(3)

                # Section content
                for line in content.split('\n'):
                    line = line.strip()
                    if not line:
                        pdf.ln(2)
                        continue

                    # Sanitize Unicode punctuation characters to safe Latin-1 equivalents
                    line = self._sanitize_pdf_text(line)

                    # Handle bullet points
                    is_bullet = bool(re.match(r'^[-\u2022*\u25aa\u25b8\u25ba\u25c6\u25cf\u25cb\u2192\u00bb·•▪◦]\s*', line))
                    if is_bullet:
                        clean_line = re.sub(r'^[-\u2022*\u25aa\u25b8\u25ba\u25c6\u25cf\u25cb\u2192\u00bb·•▪◦]+\s*', '', line)
                        line = '  ' + chr(183) + ' ' + clean_line

                    # Style headers/jobs/projects bold so they stand out
                    if not is_bullet and self._is_title_line(line, section_name):
                        # Add extra spacing above title lines to visually separate entries
                        pdf.ln(4)
                        pdf.set_font('Helvetica', 'B', 10.5)
                        pdf.set_text_color(*accent_rgb)
                        align_mode = 'L'
                    else:
                        pdf.set_font('Helvetica', '', 10)
                        pdf.set_text_color(70, 70, 70)
                        align_mode = 'J'  # Justified

                    # Encode safely
                    safe_line = line.encode('latin-1', errors='replace').decode('latin-1')
                    pdf.multi_cell(0, 5.5, safe_line, new_x="LMARGIN", new_y="NEXT", align=align_mode)

                pdf.ln(4)

            buffer = io.BytesIO()
            pdf.output(buffer)
            buffer.seek(0)
            return buffer

        except ImportError:
            return None

    # ------------------------------------------------------------------
    # Private: Styling
    # ------------------------------------------------------------------

    def _configure_styles(self, doc: Document, font_name: str, font_size: int, color: RGBColor):
        """Configure document styles for ATS compatibility."""
        # Normal style
        style = doc.styles['Normal']
        style.font.name = font_name
        style.font.size = Pt(font_size)
        style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        style.paragraph_format.space_after = Pt(2)
        style.paragraph_format.line_spacing = 1.15

        # Heading 1 (Name)
        h1 = doc.styles['Heading 1']
        h1.font.name = font_name
        h1.font.size = Pt(18)
        h1.font.bold = True
        h1.font.color.rgb = color
        h1.paragraph_format.space_before = Pt(0)
        h1.paragraph_format.space_after = Pt(4)

        # Heading 2 (Section headers)
        h2 = doc.styles['Heading 2']
        h2.font.name = font_name
        h2.font.size = Pt(13)
        h2.font.bold = True
        h2.font.color.rgb = color
        h2.paragraph_format.space_before = Pt(10)
        h2.paragraph_format.space_after = Pt(4)
        h2.paragraph_format.keep_with_next = True

    def _add_header(self, doc: Document, cv_data: Dict, font_name: str,
                    color: RGBColor, template: str):
        """Add name and contact info header."""
        name = cv_data.get('name', 'Your Name')
        contact = cv_data.get('contact', {})

        # Name
        name_para = doc.add_paragraph()
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_run = name_para.add_run(name)
        name_run.font.name = font_name
        name_run.font.size = Pt(20)
        name_run.font.bold = True
        name_run.font.color.rgb = color

        # Contact line
        contact_parts = []
        for key in ['email', 'phone', 'location', 'linkedin', 'website']:
            val = contact.get(key, '')
            if val:
                contact_parts.append(val)

        if contact_parts:
            contact_para = doc.add_paragraph()
            contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            contact_run = contact_para.add_run(' | '.join(contact_parts))
            contact_run.font.name = font_name
            contact_run.font.size = Pt(9)
            contact_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        # Separator line (subtle)
        separator = doc.add_paragraph()
        separator.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sep_run = separator.add_run('─' * 60)
        sep_run.font.size = Pt(6)
        sep_run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

    def _add_section(self, doc: Document, section_name: str, content: str,
                     font_name: str, font_size: int, color: RGBColor, template: str):
        """Add a CV section to the document."""
        # Section header
        header_text = section_name.replace('_', ' ').upper()
        heading = doc.add_heading(header_text, level=2)

        # Section content
        lines = content.split('\n')
        for line in lines:
            stripped = line.strip()
            if not stripped:
                continue

            # Detect bullet points
            is_bullet = bool(re.match(r'^[-\u2022*\u25aa\u25b8\u25ba\u25c6\u25cf\u25cb\u2192\u00bb]\s*', stripped))
            is_numbered = bool(re.match(r'^\d+[.)]\s*', stripped))

            if is_bullet or is_numbered:
                # Clean bullet prefix
                clean = re.sub(r'^[-\u2022*\u25aa\u25b8\u25ba\u25c6\u25cf\u25cb\u2192\u00bb\d.)]+\s*', '', stripped)
                para = doc.add_paragraph(style='List Bullet')
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                run = para.add_run(clean)
                run.font.name = font_name
                run.font.size = Pt(font_size)
            else:
                # Check if this looks like a job title/company/project title line
                if self._is_title_line(stripped, section_name):
                    para = doc.add_paragraph()
                    para.paragraph_format.space_before = Pt(8)
                    run = para.add_run(stripped)
                    run.font.name = font_name
                    run.font.size = Pt(font_size + 1)
                    run.font.bold = True
                    run.font.color.rgb = color
                else:
                    para = doc.add_paragraph()
                    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    run = para.add_run(stripped)
                    run.font.name = font_name
                    run.font.size = Pt(font_size)

    # ------------------------------------------------------------------
    # Private: Helpers
    # ------------------------------------------------------------------

    def _order_sections(self, sections: Dict[str, str],
                        include: List[str] = None) -> Dict[str, str]:
        """Order sections in ATS-preferred order."""
        ordered = {}

        for section_name in self.SECTION_ORDER:
            if section_name in sections:
                if include is None or section_name in include:
                    ordered[section_name] = sections[section_name]

        # Add any remaining sections not in the preferred order
        for section_name, content in sections.items():
            if section_name not in ordered and section_name not in ('header', 'contact'):
                if include is None or section_name in include:
                    ordered[section_name] = content

        return ordered

    def _is_title_line(self, line: str, section_name: str) -> bool:
        """Detect if a line is a job title, company name, or project name."""
        if section_name not in ('experience', 'education', 'projects'):
            return False

        # Lines with dates are usually title lines
        if re.search(r'(20\d{2}|19\d{2})', line):
            return True

        # Lines starting with a bullet are NOT titles
        if re.match(r'^[-•*▪▸►◆●○→»·]', line.strip()):
            return False

        # Short lines that don't end with a period look like titles
        if len(line.split()) <= 8 and not line.endswith('.'):
            return True

        # Lines ending with a colon (e.g., "My Project Name:") are titles
        if line.strip().endswith(':') and len(line.split()) <= 12:
            return True

        # Lines that are ALL CAPS or Title Case with few words
        words = line.strip().split()
        if 1 <= len(words) <= 10:
            # Title case detection (most words start with uppercase)
            upper_words = sum(1 for w in words if w[0].isupper())
            if upper_words >= len(words) * 0.6 and not line.endswith('.'):
                return True

        return False

    def _extract_contact_info(self, text: str) -> Dict[str, str]:
        """Extract contact information from text."""
        info = {}

        # Email
        email_match = re.search(r'[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}', text)
        if email_match:
            info['email'] = email_match.group()

        # Phone
        phone_match = re.search(r'[\+]?[\d\s\-\(\)]{7,15}', text)
        if phone_match:
            info['phone'] = phone_match.group().strip()

        # LinkedIn
        linkedin_match = re.search(r'linkedin\.com/in/[\w-]+', text, re.IGNORECASE)
        if linkedin_match:
            info['linkedin'] = linkedin_match.group()

        # Name (first non-empty line that doesn't look like contact info)
        lines = text.strip().split('\n')
        for line in lines[:5]:
            clean = line.strip()
            if clean and '@' not in clean and not re.match(r'^[\d\+\(\)]', clean):
                if len(clean.split()) <= 5:
                    info['name'] = clean
                    break

        return info

    # ------------------------------------------------------------------
    # Template Previews
    # ------------------------------------------------------------------

    def get_template_info(self) -> List[Dict[str, str]]:
        """Return info about available templates for the UI."""
        return [
            {
                'id': 'classic',
                'name': 'Classic Professional',
                'description': 'Traditional layout with navy blue accents. Best for corporate and finance roles.',
                'font': 'Calibri',
                'color': 'Navy Blue',
            },
            {
                'id': 'modern',
                'name': 'Modern Clean',
                'description': 'Clean, minimal layout with dark slate accents. Great for tech and creative roles.',
                'font': 'Arial',
                'color': 'Dark Slate',
            },
            {
                'id': 'minimal',
                'name': 'Elegant Minimal',
                'description': 'Sophisticated serif layout. Ideal for executive and academic positions.',
                'font': 'Garamond',
                'color': 'Dark Gray',
            },
        ]
